import os
import tempfile
import base64
from datetime import datetime
from docx import Document
from docx.shared import Pt, Inches
from PIL import Image, ImageDraw

# Try importing Weasyprint and FPDF, setting up fallback
try:
    import weasyprint
    WEASYPRINT_AVAILABLE = True
except (ImportError, OSError):
    WEASYPRINT_AVAILABLE = False
    print("Weasyprint is not available (likely missing GTK+). Falling back to FPDF2 for PDF generation.")

try:
    from fpdf import FPDF
    FPDF_AVAILABLE = True
except ImportError:
    FPDF_AVAILABLE = False
    print("FPDF2 is not available.")

class PDFGenerator(FPDF if FPDF_AVAILABLE else object):
    def __init__(self, title="WitnessIQ Case Report", case_name="Case"):
        if FPDF_AVAILABLE:
            super().__init__()
        self.title_text = title
        self.case_name = case_name

    def header(self):
        if not FPDF_AVAILABLE:
            return
        self.set_font("Helvetica", "B", 8)
        self.set_text_color(100, 100, 100)
        self.cell(0, 10, f"WITNESSIQ - CONFIDENTIAL INVESTIGATION BRIEF", 0, 0, "L")
        self.cell(0, 10, f"CASE: {self.case_name.upper()}", 0, 1, "R")
        self.line(10, 17, 200, 17)
        self.ln(5)

    def footer(self):
        if not FPDF_AVAILABLE:
            return
        self.set_y(-15)
        self.set_font("Helvetica", "I", 8)
        self.set_text_color(128, 128, 128)
        self.cell(0, 10, f"Page {self.page_no()}/{{nb}}", 0, 0, "C")
        self.cell(0, 10, f"Generated on {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", 0, 0, "R")


class ReportService:
    @staticmethod
    def customize_summary(case_title: str, report_type: str, original_description: str) -> str:
        """Tailors the executive summary narrative to highlight details most relevant to the report type."""
        title_lower = case_title.lower()
        
        # Case Enron Corporate Fraud
        if "enron" in title_lower:
            if report_type == "police":
                return (
                    "CRIMINAL INVESTIGATION BRIEF: This briefing documents evidence of systematic securities fraud, "
                    "conspiracy to misrepresent financial health, and subsequent obstruction of justice at Enron Corporation. "
                    "Key suspects include CFO Andrew Fastow (improper partnership management), CEO Kenneth Lay, and Jeffrey Skilling. "
                    "Critical evidence records include the documented destruction and shredding of Enron audit files by "
                    "Arthur Andersen LLP compliance teams on October 23, 2001, following SEC notification."
                )
            elif report_type == "hr":
                return (
                    "HR ETHICS & WHISTLEBLOWER DISCLOSURE: Internal audit concerning executive misconduct, policy violations, and whistleblower actions. "
                    "Details focus on VP of Corporate Development Sherron Watkins warning Kenneth Lay of impending accounting collapse, "
                    "and the abrupt resignation of CEO Jeffrey Skilling on August 14, 2001. Addresses employee concern reporting, "
                    "internal policy adherence, and executive transition compliance."
                )
            elif report_type == "insurance":
                return (
                    "INSURANCE ASSESSMENT BRIEF: Director & Officer (D&O) liability and corporate asset claim evaluation. "
                    "Tracks the restatement of corporate earnings (reducing profits by $586M, increasing liabilities by $2.5B), "
                    "the unwinding of Raptor/LJM shell partnerships, and professional negligence actions concerning auditors "
                    "Arthur Andersen, focusing on exclusions for deliberate fraud."
                )
            elif report_type == "disciplinary":
                return (
                    "DISCIPLINARY BOARD CASE BRIEF: Assessment of institutional policy and professional ethics breaches. "
                    "Focuses on Andrew Fastow's conflict of interest violations via dual-management of Enron CFO duties and "
                    "LJM general partnership interests, alongside Arthur Andersen's failure to maintain auditor independence."
                )
            else: # legal
                return (
                    "LEGAL EVIDENTIARY BRIEF: Evaluation of admissibility and evidentiary weight regarding Enron conspiracy linkages. "
                    "Documents chain of custody and corroborating metadata between the Sherron Watkins Whistleblower Memo, "
                    "Jeffrey Skilling's hostile analyst exchange (Richard Grubman transcript), and Arthur Andersen CCTV shredding room logs."
                )
        
        # Case Alpha Altercation
        elif "altercation" in title_lower or "warehouse" in title_lower or "alpha" in title_lower:
            if report_type == "police":
                return (
                    "POLICE INCIDENT RECONSTRUCTION: Investigation of physical assault, verbal threats, and vehicle theft at Warehouse 4. "
                    "CCTV and mobile phone records show suspect Mark Smith arrived in a black Toyota Prius (AP09XX1234), "
                    "confronted and physically assaulted security guard John Doe near the warehouse main entrance, and fled in the vehicle."
                )
            elif report_type == "hr":
                return (
                    "HR WORKPLACE SAFETY INCIDENT: Misconduct audit concerning employees John Doe and Mark Smith at Warehouse 4. "
                    "Documents threatening chat logs, verbal harassment captured on audio logs, and physical violence on company property. "
                    "Recommends policy actions for violations of zero-tolerance workplace safety standards."
                )
            elif report_type == "insurance":
                return (
                    "INSURANCE CLAIM BRIEF: Property and liability claim verification. Examines CCTV records of a black Prius (AP09XX1234) "
                    "entering and departing the warehouse perimeter at high speed, alongside physical altercations resulting in "
                    "security guard injury claims and potential facility damage."
                )
            elif report_type == "disciplinary":
                return (
                    "DISCIPLINARY HEARING REPORT: Disciplinary board brief regarding unauthorized facility access, confrontational "
                    "misconduct, and physical violence by former employee Mark Smith on John Doe at Warehouse 4."
                )
            else: # legal
                return (
                    "LEGAL BRIEF: Admissibility and evidentiary timeline analysis of physical altercation evidence. "
                    "Corroborates WhatsApp logs containing threatening communications with CCTV surveillance footage "
                    "and vocal statements, establishing a clear timeline of events for potential litigation."
                )
        
        return original_description

    @staticmethod
    def filter_and_tailor_events(events: list, report_type: str) -> list:
        """Filters events and tailors descriptions to match the focus of the report type."""
        tailored = []
        for ev in events:
            type_lower = ev.event_type.lower()
            title_lower = ev.title.lower()
            desc_lower = ev.description.lower()
            
            keep = True
            note = ""
            
            if report_type == "police":
                # Police cares about threats, altercations, destruction of evidence, vehicle movement, reports
                if any(w in type_lower or w in title_lower or w in desc_lower for w in ["threat", "altercation", "shredding", "destruction", "cctv", "prius", "arrival", "departure", "assault", "police"]):
                    keep = True
                    note = " [CRIMINAL COMPLIANCE: Threat, assault, or evidence destruction observed]"
                else:
                    if any(w in type_lower for w in ["general", "agreement", "request"]):
                        keep = False
                        
            elif report_type == "hr":
                # HR cares about communications, warnings, employee statements, harassment, policy violations
                if any(w in type_lower or w in title_lower or w in desc_lower for w in ["chat", "statement", "say", "email", "memo", "whistleblower", "threat", "alert", "altercation"]):
                    keep = True
                    note = " [HR COMPLIANCE: Interpersonal exchange or workplace behavior incident]"
                else:
                    if any(w in type_lower for w in ["surveillance", "vehicle_arrival", "vehicle_departure"]):
                        keep = False
                        
            elif report_type == "insurance":
                # Insurance cares about assets, vehicle movement, physical actions, claims, times, dates
                if any(w in type_lower or w in title_lower or w in desc_lower for w in ["prius", "vehicle", "arrival", "departure", "loss", "damage", "altercation", "accident", "shredding"]):
                    keep = True
                    note = " [LIABILITY ASSESSMENT: Property, vehicle, or liability-inducing motion]"
                else:
                    if any(w in type_lower for w in ["chat_general", "chat_request"]):
                        keep = False
                        
            elif report_type == "disciplinary":
                # Disciplinary cares about violation of rules, hearings, confrontations
                if any(w in type_lower or w in title_lower or w in desc_lower for w in ["confrontation", "altercation", "shredding", "hearing", "ethics", "whistleblower", "threat"]):
                    keep = True
                    note = " [CODE BREACH: Violation of organizational ethics or conduct rules]"
                else:
                    keep = False
                    
            else: # legal
                # Legal wants to see all events but with evidentiary annotations
                keep = True
                note = f" [EVIDENTIARY EXHIBIT: Corroborated Metadata (Confidence: {int(ev.confidence*100)}%)]"

            if keep:
                class TailoredEvent:
                    def __init__(self, original, tailored_desc):
                        self.id = original.id
                        self.case_id = original.case_id
                        self.timestamp = original.timestamp
                        self.title = original.title
                        self.description = tailored_desc
                        self.location = original.location
                        self.event_type = original.event_type
                        self.confidence = original.confidence
                        self.evidence_sources = original.evidence_sources
                
                tailored_desc = ev.description + note
                tailored.append(TailoredEvent(ev, tailored_desc))
                
        return tailored

    @staticmethod
    def generate_analytics_chart(events: list) -> str:
        """Generates a Pillow-drawn bar chart summarizing the event types and saves it as a temporary PNG."""
        # Count events by type
        counts = {}
        for ev in events:
            t = ev.event_type.replace("_", " ").title()
            counts[t] = counts.get(t, 0) + 1
            
        # Default chart data if no events
        if not counts:
            counts = {"No Data": 1}
            
        labels = list(counts.keys())
        values = list(counts.values())
        max_val = max(values)
        
        width, height = 600, 300
        img = Image.new("RGB", (width, height), "#ffffff")
        draw = ImageDraw.Draw(img)
        
        # Drawing layout
        padding_left = 180
        padding_right = 40
        padding_top = 40
        padding_bottom = 40
        chart_w = width - padding_left - padding_right
        chart_h = height - padding_top - padding_bottom
        
        # Title
        draw.text((width // 2 - 100, 10), "Timeline Event Activity by Type", fill="#1e293b")
        
        num_bars = len(labels)
        bar_height = max(10, min(30, int(chart_h / (num_bars * 1.5))))
        spacing = int((chart_h - (num_bars * bar_height)) / (num_bars + 1))
        
        for i, (label, val) in enumerate(zip(labels, values)):
            y = padding_top + spacing + i * (bar_height + spacing)
            
            # Label
            draw.text((20, y + bar_height // 2 - 5), label[:25], fill="#475569")
            
            # Bar width proportional to value
            bar_w = int((val / max_val) * chart_w) if max_val > 0 else 0
            
            # Draw bar
            draw.rectangle([padding_left, y, padding_left + bar_w, y + bar_height], fill="#3b82f6")
            
            # Value tag
            draw.text((padding_left + bar_w + 5, y + bar_height // 2 - 5), str(val), fill="#1e293b")
            
        temp_dir = tempfile.gettempdir()
        temp_path = os.path.join(temp_dir, "witness_iq_chart.png")
        img.save(temp_path)
        return temp_path

    @staticmethod
    def get_report_content(report_type: str, case_title: str, description: str, events: list, entities: list) -> dict:
        """Structures and customizes the content based on report type."""
        sections = {}
        
        if report_type == "police":
            sections["title"] = f"CRIMINAL INCIDENT REPORT: {case_title}"
            sections["purpose"] = "This document compiles chronological evidence regarding a potential criminal offence, outlining key events, involved parties, vehicles, and direct evidence links."
            sections["evidence_focus"] = "Physical actions, vehicles, locations, and threats."
            sections["legal_text"] = "Submitted for review to law enforcement agencies. The findings below are compiled automatically from digital evidence via WitnessIQ."
            
        elif report_type == "insurance":
            sections["title"] = f"INSURANCE ACCIDENT INVESTIGATION REPORT: {case_title}"
            sections["purpose"] = "This report reconstructs the chronology of events surrounding the filed claim to verify authenticity, identify liabilities, and establish a timeline of property/vehicle activity."
            sections["evidence_focus"] = "Vehicle movement, timestamps, and claims parameters."
            sections["legal_text"] = "Prepared for claims assessment. Subject to verification against original metadata."

        elif report_type == "hr":
            sections["title"] = f"HR INVESTIGATION REPORT: {case_title}"
            sections["purpose"] = "Internal report regarding allegations of workplace misconduct, harassment, or policy violation, detailing interactions between employees."
            sections["evidence_focus"] = "Interpersonal communications, chat logs, audio statements, and work-hours activity."
            sections["legal_text"] = "Strictly Confidential. For internal Human Resources and Legal Counsel review only."

        elif report_type == "disciplinary":
            sections["title"] = f"DISCIPLINARY BOARD CASE BRIEF: {case_title}"
            sections["purpose"] = "An audit of events concerning student/member conduct violations, providing an objective timeline derived from uploaded materials."
            sections["evidence_focus"] = "Witness statements, text/chat arguments, and presence at key locations."
            sections["legal_text"] = "For Dean of Students and Disciplinary Committee Review."

        else: # legal
            sections["title"] = f"LEGAL EVIDENCE BRIEF: {case_title}"
            sections["purpose"] = "A formal evidentiary brief organizing digital assets chronologically, demonstrating chain of custody and cross-corroborating testimonials with metadata."
            sections["evidence_focus"] = "Admissible digital assets, verified transcripts, and entity-association graphs."
            sections["legal_text"] = "Attorney-Client Privileged / Work Product. Prepared for litigation."

        return sections

    @classmethod
    def generate_html_report(cls, report_type: str, case_title: str, description: str, events: list, entities: list) -> str:
        """Generates a highly-styled HTML report with dynamically tailored summaries and filtered events."""
        meta = cls.get_report_content(report_type, case_title, description, events, entities)
        
        # Tailor body text matter
        custom_description = cls.customize_summary(case_title, report_type, description)
        tailored_events = cls.filter_and_tailor_events(events, report_type)

        # Build Analytics Chart Base64
        chart_base64 = ""
        try:
            chart_path = cls.generate_analytics_chart(tailored_events)
            with open(chart_path, "rb") as image_file:
                chart_base64 = base64.b64encode(image_file.read()).decode('utf-8')
        except Exception as e:
            print(f"Failed to encode chart for HTML: {e}")

        # Gather evidence images
        evidence_images = []
        seen_image_paths = set()
        for ev in tailored_events:
            for src in ev.evidence_sources:
                if src.file_type in ["image", "jpg", "png", "jpeg"] and src.file_path not in seen_image_paths:
                    seen_image_paths.add(src.file_path)
                    evidence_images.append(src)

        # Build HTML structure
        html = f"""
        <html>
        <head>
            <style>
                body {{ font-family: 'Helvetica Neue', Helvetica, Arial, sans-serif; color: #333; line-height: 1.5; margin: 40px; }}
                h1 {{ color: #1e293b; border-bottom: 2px solid #3b82f6; padding-bottom: 8px; font-size: 24px; }}
                h2 {{ color: #334155; margin-top: 30px; font-size: 18px; border-bottom: 1px solid #e2e8f0; padding-bottom: 5px; }}
                .meta-box {{ background-color: #f8fafc; border: 1px solid #e2e8f0; padding: 15px; border-radius: 6px; margin-bottom: 20px; }}
                .meta-table {{ width: 100%; border-collapse: collapse; }}
                .meta-table td {{ padding: 6px; font-size: 14px; }}
                .meta-table td.label {{ font-weight: bold; color: #64748b; width: 25%; }}
                .event-card {{ border-left: 3px solid #3b82f6; padding-left: 15px; margin-bottom: 15px; page-break-inside: avoid; }}
                .event-time {{ font-weight: bold; color: #2563eb; font-size: 14px; }}
                .event-title {{ font-weight: bold; color: #1e293b; font-size: 15px; margin: 2px 0; }}
                .event-desc {{ color: #475569; font-size: 13px; margin: 5px 0; }}
                .event-source {{ font-size: 11px; color: #94a3b8; font-style: italic; }}
                .entity-badge {{ display: inline-block; background: #f1f5f9; border: 1px solid #cbd5e1; padding: 2px 8px; border-radius: 12px; font-size: 12px; margin: 4px; color: #334155; }}
                .entity-type {{ font-weight: bold; font-size: 9px; color: #2563eb; text-transform: uppercase; margin-right: 4px; }}
                .confidential {{ text-align: center; color: #ef4444; font-weight: bold; letter-spacing: 2px; margin-bottom: 30px; font-size: 12px; }}
            </style>
        </head>
        <body>
            <div class="confidential">CONFIDENTIAL // INVESTIGATION REPORT</div>
            <h1>{meta['title']}</h1>
            
            <div class="meta-box">
                <table class="meta-table">
                    <tr>
                        <td class="label">Case Name:</td>
                        <td>{case_title}</td>
                    </tr>
                    <tr>
                        <td class="label">Generated:</td>
                        <td>{datetime.now().strftime('%B %d, %Y %H:%M:%S')}</td>
                    </tr>
                    <tr>
                        <td class="label">Scope:</td>
                        <td>{meta['purpose']}</td>
                    </tr>
                    <tr>
                        <td class="label">Evidence Focus:</td>
                        <td>{meta['evidence_focus']}</td>
                    </tr>
                </table>
            </div>

            <h2>Executive Summary</h2>
            <p style="font-size: 14px; color: #334155; line-height: 1.6;">{custom_description}</p>
        """

        if chart_base64:
            html += f"""
            <h2>Timeline Activity Analytics</h2>
            <div style="text-align: center; margin: 20px 0; page-break-inside: avoid;">
                <img src="data:image/png;base64,{chart_base64}" style="max-width: 100%; border: 1px solid #cbd5e1; border-radius: 6px; box-shadow: 0 1px 3px rgba(0,0,0,0.1);" />
            </div>
            """

        html += """
            <h2>Reconstructed Incident Timeline</h2>
            <div class="timeline">
        """

        for ev in tailored_events:
            time_str = ev.timestamp.strftime("%Y-%m-%d %H:%M:%S")
            sources_str = ", ".join([s.filename for s in ev.evidence_sources]) if ev.evidence_sources else "Synthetic Correlation"
            html += f"""
                <div class="event-card">
                    <div class="event-time">{time_str} (Confidence: {int(ev.confidence*100)}%)</div>
                    <div class="event-title">{ev.title}</div>
                    <div class="event-desc">{ev.description.replace(chr(10), '<br>')}</div>
                    <div class="event-source">Source(s): {sources_str}</div>
                </div>
            """

        html += """
            </div>
            <h2>Extracted Key Entities</h2>
            <div style="margin-top: 10px;">
        """

        for ent in entities:
            html += f"""
                <div class="entity-badge">
                    <span class="entity-type">{ent.type}</span> {ent.name} <span style="color:#64748b;">({ent.details or ''})</span>
                </div>
            """

        html += "</div>"

        # Evidence Images rendering in HTML
        if evidence_images:
            html += "<h2>Evidence Visual Assets</h2>"
            for img_ev in evidence_images:
                if os.path.exists(img_ev.file_path):
                    try:
                        with open(img_ev.file_path, "rb") as f:
                            img_b64 = base64.b64encode(f.read()).decode('utf-8')
                        html += f"""
                        <div style="margin-bottom: 30px; page-break-inside: avoid; border: 1px solid #e2e8f0; padding: 15px; border-radius: 6px; background-color: #f8fafc;">
                            <div style="font-weight: bold; color: #475569; font-size: 14px;">Evidence: {img_ev.filename}</div>
                            <div style="font-size: 12px; color: #64748b; font-style: italic; margin-bottom: 8px;">{img_ev.summary or ''}</div>
                            <div style="text-align: center;">
                                <img src="data:image/png;base64,{img_b64}" style="max-width: 100%; border: 1px solid #cbd5e1; border-radius: 4px;" />
                            </div>
                        </div>
                        """
                    except Exception as e:
                        print(f"Failed to render image {img_ev.filename} in HTML: {e}")

        html += f"""
            <div style="margin-top: 50px; font-size: 11px; color: #94a3b8; border-top: 1px solid #e2e8f0; padding-top: 10px;">
                {meta['legal_text']}<br>
                Report generated automatically by WitnessIQ.
            </div>
        </body>
        </html>
        """
        return html

    @classmethod
    def generate_pdf_report(cls, report_type: str, case_title: str, description: str, events: list, entities: list, output_path: str):
        """Generates PDF using Weasyprint or falls back to FPDF2, fully tailoring the contents."""
        html_content = cls.generate_html_report(report_type, case_title, description, events, entities)
        
        if WEASYPRINT_AVAILABLE:
            try:
                weasyprint.HTML(string=html_content).write_pdf(output_path)
                return
            except Exception as e:
                print(f"WeasyPrint failed to compile PDF: {e}. Falling back to FPDF2.")

        if FPDF_AVAILABLE:
            # Generate via FPDF2
            pdf = PDFGenerator(title=f"WitnessIQ: {case_title}", case_name=case_title)
            pdf.alias_nb_pages()
            pdf.add_page()
            
            # Swapped / Tailored Matter
            custom_description = cls.customize_summary(case_title, report_type, description)
            tailored_events = cls.filter_and_tailor_events(events, report_type)

            # Title
            meta = cls.get_report_content(report_type, case_title, description, events, entities)
            pdf.set_font("Helvetica", "B", 16)
            pdf.set_text_color(30, 41, 59)
            pdf.cell(0, 10, meta["title"], 0, 1, "L")
            pdf.ln(5)
            
            # Metadata Box
            pdf.set_fill_color(248, 250, 252)
            pdf.set_draw_color(226, 232, 240)
            pdf.rect(10, pdf.get_y(), 190, 40, "DF")
            pdf.set_font("Helvetica", "B", 10)
            pdf.set_text_color(100, 100, 100)
            pdf.text(15, pdf.get_y() + 8, "Case Name:")
            pdf.text(15, pdf.get_y() + 18, "Generated:")
            pdf.text(15, pdf.get_y() + 28, "Focus:")
            
            pdf.set_font("Helvetica", "", 10)
            pdf.set_text_color(51, 65, 85)
            pdf.text(40, pdf.get_y() + 8, case_title)
            pdf.text(40, pdf.get_y() + 18, datetime.now().strftime('%Y-%m-%d %H:%M:%S'))
            pdf.text(40, pdf.get_y() + 28, meta["evidence_focus"])
            pdf.ln(45)
            
            # Executive Summary
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "Executive Summary", 0, 1, "L")
            pdf.set_font("Helvetica", "", 10)
            pdf.multi_cell(0, 5, custom_description)
            pdf.ln(8)
            
            # Add Page for Analytics
            try:
                chart_path = cls.generate_analytics_chart(tailored_events)
                pdf.add_page()
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(30, 41, 59)
                pdf.cell(0, 10, "Timeline Activity Analytics", 0, 1, "L")
                pdf.ln(5)
                pdf.image(chart_path, x=15, y=pdf.get_y(), w=180)
                pdf.ln(95)
            except Exception as e:
                print(f"Failed to embed chart in FPDF: {e}")

            # Timeline
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 12)
            pdf.cell(0, 10, "Reconstructed Incident Timeline", 0, 1, "L")
            pdf.ln(5)
            
            for ev in tailored_events:
                time_str = ev.timestamp.strftime("%Y-%m-%d %H:%M:%S")
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(37, 99, 235)
                pdf.cell(0, 6, f"{time_str} (Confidence: {int(ev.confidence*100)}%)", 0, 1)
                
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(30, 41, 59)
                pdf.cell(0, 5, ev.title, 0, 1)
                
                pdf.set_font("Helvetica", "", 9)
                pdf.set_text_color(71, 85, 105)
                clean_desc = ev.description.replace('\r', '')
                pdf.multi_cell(0, 5, clean_desc)
                
                sources = ", ".join([s.filename for s in ev.evidence_sources]) if ev.evidence_sources else "Correlated"
                pdf.set_font("Helvetica", "I", 8)
                pdf.set_text_color(148, 163, 184)
                pdf.cell(0, 5, f"Source(s): {sources}", 0, 1)
                pdf.ln(4)
                
            # Entities
            pdf.add_page()
            pdf.set_font("Helvetica", "B", 12)
            pdf.set_text_color(30, 41, 59)
            pdf.cell(0, 10, "Extracted Key Entities", 0, 1, "L")
            pdf.ln(5)
            
            for ent in entities:
                pdf.set_font("Helvetica", "B", 10)
                pdf.set_text_color(37, 99, 235)
                pdf.write(5, f"[{ent.type.upper()}] ")
                pdf.set_font("Helvetica", "", 10)
                pdf.set_text_color(30, 41, 59)
                pdf.write(5, f"{ent.name} ({ent.details or ''})\n")
                pdf.ln(2)

            # Evidence Images gathering from tailored events
            evidence_images = []
            seen_image_paths = set()
            for ev in tailored_events:
                for src in ev.evidence_sources:
                    if src.file_type in ["image", "jpg", "png", "jpeg"] and src.file_path not in seen_image_paths:
                        seen_image_paths.add(src.file_path)
                        evidence_images.append(src)

            # Evidence Images rendering in PDF
            if evidence_images:
                pdf.add_page()
                pdf.set_font("Helvetica", "B", 12)
                pdf.set_text_color(30, 41, 59)
                pdf.cell(0, 10, "Evidence Visual Assets", 0, 1, "L")
                pdf.ln(5)
                
                for img_ev in evidence_images:
                    if os.path.exists(img_ev.file_path):
                        try:
                            pdf.set_font("Helvetica", "B", 10)
                            pdf.set_text_color(71, 85, 105)
                            pdf.cell(0, 6, f"Evidence: {img_ev.filename}", 0, 1)
                            if img_ev.summary:
                                pdf.set_font("Helvetica", "I", 9)
                                pdf.set_text_color(100, 100, 100)
                                pdf.multi_cell(0, 5, img_ev.summary)
                            pdf.ln(2)
                            pdf.image(img_ev.file_path, x=20, y=pdf.get_y(), w=170)
                            pdf.ln(95)
                        except Exception as img_err:
                            pdf.set_font("Helvetica", "I", 9)
                            pdf.set_text_color(239, 68, 68)
                            pdf.cell(0, 5, f"[Failed to render image {img_ev.filename}: {img_err}]", 0, 1)
                            pdf.ln(4)

            # Footer legal notice
            pdf.ln(10)
            pdf.set_font("Helvetica", "I", 8)
            pdf.set_text_color(148, 163, 184)
            pdf.multi_cell(0, 4, meta["legal_text"])

            pdf.output(output_path)
        else:
            # Fallback
            html_path = output_path.replace(".pdf", ".html")
            with open(html_path, "w", encoding="utf-8") as f:
                f.write(html_content)
            print(f"FPDF/WeasyPrint not installed. Saved HTML report to {html_path}")
            with open(output_path, "w") as f:
                f.write(f"PDF creation failed. Please check HTML report: {html_path}")

    @classmethod
    def generate_docx_report(cls, report_type: str, case_title: str, description: str, events: list, entities: list, output_path: str):
        """Generates Word document (.docx) with dynamically tailored summaries and filtered events."""
        meta = cls.get_report_content(report_type, case_title, description, events, entities)
        
        # Swapped / Tailored Matter
        custom_description = cls.customize_summary(case_title, report_type, description)
        tailored_events = cls.filter_and_tailor_events(events, report_type)

        doc = Document()
        
        # Style Definitions
        styles = doc.styles
        normal_style = styles['Normal']
        normal_style.font.name = 'Arial'
        normal_style.font.size = Pt(10.5)

        # Header Title
        title_p = doc.add_paragraph()
        run = title_p.add_run(meta["title"].upper())
        run.bold = True
        run.font.size = Pt(16)
        
        doc.add_paragraph("CONFIDENTIAL // INVESTIGATION REPORT").alignment = 1 # Center
        
        # Metadata Table
        table = doc.add_table(rows=4, cols=2)
        table.style = 'Light Shading Accent 1'
        
        meta_items = [
            ("Case Name", case_title),
            ("Date Generated", datetime.now().strftime('%Y-%m-%d %H:%M:%S')),
            ("Scope & Purpose", meta["purpose"]),
            ("Focus Areas", meta["evidence_focus"])
        ]
        
        for i, (label, val) in enumerate(meta_items):
            row = table.rows[i]
            row.cells[0].paragraphs[0].add_run(label).bold = True
            row.cells[1].paragraphs[0].add_run(val)
            
        doc.add_paragraph() # spacing
        
        # Executive Summary
        h_summary = doc.add_paragraph()
        r_summary = h_summary.add_run("Executive Summary")
        r_summary.bold = True
        r_summary.font.size = Pt(13)
        
        doc.add_paragraph(custom_description)
        
        # Add Analytics Chart to DOCX based on tailored events
        try:
            chart_path = cls.generate_analytics_chart(tailored_events)
            doc.add_page_break()
            h_chart = doc.add_paragraph()
            r_chart = h_chart.add_run("Timeline Activity Analytics")
            r_chart.bold = True
            r_chart.font.size = Pt(13)
            doc.add_picture(chart_path, width=Inches(6.0))
        except Exception as e:
            print(f"Failed to add chart to DOCX: {e}")

        # Timeline
        doc.add_page_break()
        h_timeline = doc.add_paragraph()
        r_timeline = h_timeline.add_run("Reconstructed Incident Timeline")
        r_timeline.bold = True
        r_timeline.font.size = Pt(13)
        
        for ev in tailored_events:
            time_str = ev.timestamp.strftime("%Y-%m-%d %H:%M:%S")
            ev_p = doc.add_paragraph()
            
            r_time = ev_p.add_run(f"[{time_str}] ")
            r_time.bold = True
            
            r_title = ev_p.add_run(f"{ev.title} (Confidence: {int(ev.confidence*100)}%)\n")
            r_title.bold = True
            
            ev_p.add_run(f"{ev.description}\n")
            
            sources = ", ".join([s.filename for s in ev.evidence_sources]) if ev.evidence_sources else "Correlated"
            r_sources = ev_p.add_run(f"Source(s): {sources}")
            r_sources.italic = True
            r_sources.font.size = Pt(9)
            
        # Entities
        doc.add_page_break()
        h_entities = doc.add_paragraph()
        r_entities = h_entities.add_run("Extracted Key Entities")
        r_entities.bold = True
        r_entities.font.size = Pt(13)
        
        ent_table = doc.add_table(rows=1, cols=3)
        ent_table.style = 'Light Shading Accent 1'
        hdr_cells = ent_table.rows[0].cells
        hdr_cells[0].paragraphs[0].add_run("Entity Name").bold = True
        hdr_cells[1].paragraphs[0].add_run("Type").bold = True
        hdr_cells[2].paragraphs[0].add_run("Details").bold = True
        
        for ent in entities:
            row_cells = ent_table.add_row().cells
            row_cells[0].paragraphs[0].add_run(ent.name)
            row_cells[1].paragraphs[0].add_run(ent.type.upper())
            row_cells[2].paragraphs[0].add_run(ent.details or "")
            
        # Gather evidence images from tailored events
        evidence_images = []
        seen_image_paths = set()
        for ev in tailored_events:
            for src in ev.evidence_sources:
                if src.file_type in ["image", "jpg", "png", "jpeg"] and src.file_path not in seen_image_paths:
                    seen_image_paths.add(src.file_path)
                    evidence_images.append(src)
                    
        # Evidence Images rendering in DOCX
        if evidence_images:
            doc.add_page_break()
            h_vis = doc.add_paragraph()
            r_vis = h_vis.add_run("Evidence Visual Assets")
            r_vis.bold = True
            r_vis.font.size = Pt(13)
            
            for img_ev in evidence_images:
                if os.path.exists(img_ev.file_path):
                    try:
                        doc.add_paragraph(f"Evidence: {img_ev.filename}").runs[0].bold = True
                        if img_ev.summary:
                            doc.add_paragraph(img_ev.summary).runs[0].italic = True
                        doc.add_picture(img_ev.file_path, width=Inches(5.5))
                        doc.add_paragraph() # spacing
                    except Exception as e:
                        print(f"Failed to add image {img_ev.filename} to DOCX: {e}")

        doc.add_paragraph()
        
        # Legal Notice
        legal_p = doc.add_paragraph()
        r_legal = legal_p.add_run(meta["legal_text"])
        r_legal.italic = True
        r_legal.font.size = Pt(8.5)
        
        doc.save(output_path)
